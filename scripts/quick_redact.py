#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
法律文件脱敏处理 - 一键脱敏脚本（v2）

完整复刻 HTML 应用 identifyRedactions() 的识别逻辑，
确保与 HTML 脱敏效果一致。

用法：python quick_redact.py <docx文件路径>
"""

import os
import sys
import json
import re
from datetime import datetime
from pathlib import Path

try:
    from docx import Document
except ImportError:
    print("错误: 缺少 python-docx 库，请运行: pip install python-docx")
    sys.exit(1)


# ============================================================
# 规则定义（从 HTML DEFAULT_RULES 完整复刻）
# ============================================================

RULES = [
    {
        'category': 'email',
        'patterns': [
            r'[\w\-\.\]+@[\w\-]+(?:\.[\w\-]+)+',
            r'邮箱\s*[：:]?\s*([\w\-\.\]+@[\w\-]+(?:\.[\w\-]+)+)',
        ],
        'replacement': '【邮箱${index}】',
        'useCaptureGroup': True,
    },
    {
        'category': 'id_card',
        'patterns': [
            r'[1-9]\d{5}(18|19|20)\d{2}(0[1-9]|1[0-2])(0[1-9]|[12]\d|3[01])\d{3}[0-9Xx]',
        ],
        'replacement': '【身份证号${index}】',
    },
    {
        'category': 'organization',
        'patterns': [
            r'([\u4e00-\u9fa5]{2,20})(有限公司|股份公司|集团)',
            r'([A-Za-z][A-Za-z0-9\s&]{2,30})(?:Company|Co\.|Corporation|Corp\.|Ltd\.|Limited|LLC|Inc\.|GmbH|S\.A\.|S\.A|Sp\. z o\.o\.|K\.K\.)\b',
            r'([\u4e00-\u9fa5]{2,10})(?:有限责任公司|股份有限公司)',
        ],
        'replacement': '【公司${index}】',
    },
    {
        'category': 'date',
        'patterns': [
            r'\d{4}\s{0,2}年\s{0,2}\d{1,2}\s{0,2}月\s{0,2}\d{1,2}\s{0,2}日',
            r'\d{4}-\d{1,2}-\d{1,2}',
            r'\d{4}/\d{1,2}/\d{1,2}',
            r'\d{1,2}月\d{1,2}日',
        ],
        'replacement': '【日期${index}】',
    },
    {
        'category': 'price',
        'patterns': [
            r'人民币\s*(\d+(?:,\d{3})*(?:\.\d+)?)\s*[元万元]',
            r'￥\s*(\d+(?:,\d{3})*(?:\.\d+)?)\s*元',
            r'\d+(?:,\d{3})*(?:\.\d+)?[万千万元]',
            r'(?:人民币|￥)\s*([壹贰叁肆伍陆柒捌玖零拾佰仟万亿千百十]+元整?)',
            r'(?:人民币|￥)\s*([壹贰叁肆伍陆柒捌玖零拾佰仟万亿千百十]+[角分])',
            r'[壹贰叁肆伍陆柒捌玖零拾佰仟万亿千百十]+元整?',
            r'[壹贰叁肆伍陆柒捌玖零拾佰仟万亿千百十]+[角分]',
            r'(?:RMB|USD|CNY|HKD|JPY|EUR|GBP)\s+(\d+(?:,\d{3})*(?:\.\d+)?)',
            r'(\d+(?:,\d{3})*(?:\.\d+)?)\s*[～\－－至]\s*(\d+(?:,\d{3})*(?:\.\d+)?)',
            r'(?:RMB|USD|CNY|HKD|JPY|EUR|GBP)\s+\d+(?:,\d{3})*(?:\.\d+)?\s*[～\－－至]\s*(\d+(?:,\d{3})*(?:\.\d+)?)',
        ],
        'replacement': '【金额${index}】',
        'useCaptureGroup': True,
    },
    {
        'category': 'person_name',
        'patterns': [
            r'(?:甲方|乙方|丙方|丁方|戊方|己方|原告方|被告方|委托方|受托方|发包方|承包方|采购方|供应商|卖方|买方|出租方|承租方|转让方|受让方|投资方|被投资方|借款人|贷款人|保证人|抵押人|出质人|收款人|付款人|债务人|债权人|申请人|被申请人|法定代表人|联系人)\s*[：:]\s*([\u4e00-\u9fa5]{2,20})',
        ],
        'replacement': '【人员${index}】',
        'useCaptureGroup': True,
    },
    {
        'category': 'phone',
        'patterns': [
            r'联系电话\s*[：:]\s*(1[3-9]\d{9})',
            r'联系电话\s*[：:]\s*(0\d{2,3}-?\d{7,8})',
            r'联系电话\s*[：:]\s*\((0\d{2,3})\)\d{7,8}',
            r'电话\s*[：:]\s*(1[3-9]\d{9})',
            r'电话\s*[：:]\s*(0\d{2,3}-?\d{7,8})',
            r'手机\s*[：:]\s*(1[3-9]\d{9})',
        ],
        'useCaptureGroup': True,
        'replacement': '【电话${index}】',
    },
    {
        'category': 'credit_code',
        'patterns': [
            r'统一社会信用代码\s*[：:]\s*([0-9A-HJ-NPQRTUWXY]{15,18})',
            r'统一社会信用代码\s*[：:]\s*([0-9a-hj-npqrtuwxy]{15,18})',
        ],
        'useCaptureGroup': True,
        'replacement': '【代码-统一社会信用代码${index}】',
    },
    {
        'category': 'bank_account',
        'patterns': [
            r'账号\s*[：:]\s*(\d{16,19})',
            r'银行账号\s*[：:]\s*(\d{16,19})',
            r'账户\s*[：:]\s*(\d{16,19})',
            r'\d{4}\s+\d{4}\s+\d{4}\s+\d{4}(?:\s+\d{4})?',
        ],
        'replacement': '【账号${index}】',
        'useCaptureGroup': True,
    },
    {
        'category': 'org_code',
        'patterns': [
            r'组织机构代码\s*[：:]\s*(\d{8}-?[0-9X])',
        ],
        'useCaptureGroup': True,
        'replacement': '【代码-组织机构代码${index}】',
    },
    {
        'category': 'patent_code',
        'patterns': [
            r'(?:专利|商标|著作权)申请号\s*[：:]\s*([A-Z0-9\.]+)',
            r'(?:JP|US|EP|CN|WO|KR|GB|DE|FR|CA|AU)(?:\d{4,}|\d{4}/\d{6})[A-Z]\d?',
            r'[A-Z]{2}\d{6,}[A-Z]{0,2}\d?',
        ],
        'useCaptureGroup': True,
        'replacement': '【代码-专利申请号${index}】',
    },
    {
        'category': 'file_code',
        'patterns': [
            r'合同编号\s*[：:]\s*([A-Z0-9\-]+)',
            r'文件编号\s*[：:]\s*([A-Z0-9\-]+)',
            r'文件号\s*[：:]\s*([A-Z0-9\-]+)',
        ],
        'useCaptureGroup': True,
        'replacement': '【代码-文件编号${index}】',
    },
    {
        'category': 'case_number',
        'patterns': [
            r'([（(]\d{4}[)）][^民刑行知赔]{1,6}?(?:民|刑|行|知|赔)[^初终再监执字第]{0,5}?\d+[号])(?:[，。、,;．.？])?',
            r'(\(\d{4}\)[^\"]{2,8}?\d+号)(?:[，。、,;．.？])?',
            r'(（\d{4}）[^\"]{2,8}?\d+号)(?:[，。、,;．.？])?',
            r'(\d{4}年（[A-Z]{1,5}）第\d+号)(?:[，。、,;．.？])?',
            r'(Case No\.(?:\s+[\d:]+-cv-)?\d+)(?:[，。、,;．.？])?',
            r'((?:Civil Action) No\.\s+[\d:]+-cv-\d+)(?:[，。、,;．.？])?',
            r'(\[\d{4}\]\s+[A-Z]{1,4}\s+\d+\s+\([A-Z]+\))(?:[，。、,;．.？])?',
            r'(Az\.:\s+\d+[A-Z]?\s+\d+/\d{2,4})(?:[，。、,;．.？])?',
            r'(RG\s+\d+/\d+)(?:[，。、,;．.？])?',
        ],
        'useCaptureGroup': True,
        'replacement': '【案号${index}】',
    },
    {
        'category': 'project_name',
        'patterns': [
            r'([\u4e00-\u9fa5]{2,10})(项目|工程|系统|平台|计划)',
            r'([\u4e00-\u9fa5]{2,10})(研发|建设|实施)项目',
        ],
        'replacement': '【项目${index}】',
    },
    {
        'category': 'address',
        'patterns': [
            r'(?:地址|住所|住址|注册地址|办公地址|住所地)\s*[：:]\s*([\u4e00-\u9fa5\w\d\u3000]+)',
            r'([\u4e00-\u9fa5]{2,4}市[\u4e00-\u9fa5]{1,4}区[\u4e00-\u9fa5\d]+(?:路|街|道|大道)(?:\d+|[\u4e00-\u9fa5]+)楼\d{1,4}-\d{1,4})',
        ],
        'useCaptureGroup': True,
        'replacement': '【地址${index}】',
    },
    {
        'category': 'sensitive',
        'patterns': [],
        'replacement': '【敏感信息${index}】',
    },
]

# 排除词集合（person_name 类别）
EXCLUDED_TERMS = {
    '甲方', '乙方', '丙方', '丁方', '戊方', '己方',
    '原告方', '被告方', '委托方', '受托方', '发包方', '承包方',
    '采购方', '供应商', '卖方', '买方',
    '出租方', '承租方', '转让方', '受让方', '投资方', '被投资方',
    '借款人', '贷款人', '保证人', '抵押人', '出质人',
    '收款人', '付款人', '债务人', '债权人',
    '申请人', '被申请人',
    '签署时间', '签署日期', '送达地址', '住址', '通讯地址',
    '联系方式', '联系电话', '联系邮箱', '授权代表', '法定代表人',
    '证件号码', '身份证号', '统一社会信用代码', '注册地址', '办公地址',
    '邮政编码', '开户行', '账号',
}

# 排除项目名称集合（project_name 类别）
EXCLUDED_PROJECTS = {
    '股权激励计划', '公司高管股权激励计划', '员工股权激励计划',
    '限制性股票激励计划', '股票期权激励计划',
    '股权激励', '员工持股计划', '限制性股票', '股票期权',
    '增资计划', '增资项目', '融资计划', '融资项目',
    '投资计划', '投资项目', '合作计划', '合作项目',
    '全国项目', '国家项目', '省级项目', '市级项目',
    '专项资金', '财政资金', '预算资金', '自有资金',
    '研发项目', '建设项目', '实施项目',
    '工程项目', '系统项目', '平台项目',
}

# 类型名称映射
TYPE_NAMES = {
    'id_card': '身份证号',
    'organization': '组织机构',
    'date': '日期',
    'price': '金额',
    'person_name': '人名',
    'phone': '联系电话',
    'email': '联系邮箱',
    'credit_code': '代码-统一社会信用代码',
    'bank_account': '账号',
    'org_code': '代码-组织机构代码',
    'patent_code': '代码-专利申请号',
    'file_code': '代码-文件编号',
    'case_number': '案号',
    'project_name': '项目名称',
    'address': '地址',
    'blacklist': '黑名单',
    'sensitive': '敏感信息',
}


# ============================================================
# 核心识别逻辑（完整复刻 HTML identifyRedactions）
# ============================================================

def is_covered_by_whitelist(original, whitelist):
    """白名单子串覆盖检查：匹配内容是否包含白名单项"""
    for wl in whitelist:
        if wl in original:
            return True
    return False


def should_skip(original, match, full_text, rule, whitelist, identified_originals):
    """
    复刻 HTML 中对每个匹配的全部后处理过滤器。
    返回 True 表示应跳过该匹配。

    注意：去重和子串判断基于 original（捕获组文本），与 HTML 的
    isDuplicate() 和子串检查逻辑完全一致。
    """
    category = rule.get('category', '')
    start_index = match.start()
    match_text = match.group(0)

    # 1. price 边界检查（基于 match[0] 的位置）
    if category == 'price':
        match_end_index = start_index + len(match_text)
        if start_index > 0 and full_text[start_index - 1].isdigit():
            return True
        if match_end_index < len(full_text):
            next_char = full_text[match_end_index]
            if next_char.isdigit() or next_char == '-':
                return True

    # 2. organization 边界检查（等价于 HTML regex 中的 lookbehind
    #    (?<![\u4e00-\u9fa5\d])，防止 "2025年上海这个那个有限公司" 中
    #    "年"后面的匹配被误识别。注意：project_name 不需要此检查。）
    if category == 'organization' and start_index > 0:
        char_before = full_text[start_index - 1]
        if re.match(r'[\u4e00-\u9fa5\d]', char_before):
            return True

    # 3. 跨段匹配排除
    if category in ('person_name', 'address') and '\n' in match_text:
        return True

    # 4. 已脱敏格式检查
    if re.match(r'^【.+】$', original):
        return True

    # 5. EXCLUDED_TERMS（person_name）
    if category == 'person_name' and original in EXCLUDED_TERMS:
        return True

    # 6. EXCLUDED_PROJECTS（project_name）
    if category == 'project_name' and original in EXCLUDED_PROJECTS:
        return True

    # 7. 地址过滤
    if category == 'address':
        if '联系电话' in original:
            return True
        if len(original) < 3:
            return True

    # 8. 白名单子串覆盖
    if is_covered_by_whitelist(original, whitelist):
        return True

    # 9. 同类别子串覆盖（基于 original，与 HTML 一致：
    #    this.redactions.some(r => r.type === rule.category && r.original !== original && r.original.includes(original)))
    for existing_original, existing_type in identified_originals:
        if existing_type == category and existing_original != original and original in existing_original:
            return True

    # 9. 去重（基于 original，与 HTML isDuplicate 一致：
    #    this.redactions.some(r => r.original === original)）
    if any(o == original for o, _ in identified_originals):
        return True

    return False


def identify_redactions(full_text, disabled_rules, whitelist, blacklist, custom_types):
    """
    完整复刻 HTML identifyRedactions() 的识别逻辑。

    Args:
        full_text: 全文（段落用 \\n 连接）
        disabled_rules: set of disabled category names
        whitelist: set of whitelist strings
        blacklist: list of {original, type} dicts
        custom_types: dict of {type_name: [{original, id}, ...]}

    Returns:
        list of {original, replacement, type, location} dicts
    """
    redactions = []
    type_counters = {}
    identified_originals = []  # 存储 (original, type) 元组，与 HTML redactions 列表一致

    # ---- 1. 自动规则匹配 ----
    for rule in RULES:
        category = rule.get('category', '')
        if category in disabled_rules:
            continue

        for pattern in rule.get('patterns', []):
            regex = re.compile(pattern)
            for match in regex.finditer(full_text):
                # 提取 original（捕获组逻辑）
                use_cg = rule.get('useCaptureGroup', False)
                if use_cg and match.lastindex and match.lastindex >= 1:
                    original = match.group(1)
                else:
                    original = match.group(0)

                # 过滤
                if should_skip(original, match, full_text, rule, whitelist, identified_originals):
                    continue

                # 计数 + 替换
                type_counters[category] = type_counters.get(category, 0) + 1
                replacement = rule.get('replacement', '【脱敏内容】').replace('${index}', str(type_counters[category]))

                redactions.append({
                    'original': original,
                    'replacement': replacement,
                    'type': category,
                    'location': '',
                })
                identified_originals.append((original, category))

                # 电话二次扫描
                if category == 'phone':
                    full_match_text = match.group(0)
                    if any(kw in full_match_text for kw in ('联系电话', '电话：', '手机：', '电话:')):
                        # 找到当前行末尾
                        line_end = full_text.find('\n', match.end())
                        if line_end == -1:
                            line_end = len(full_text)
                        remaining = full_text[match.end():line_end]

                        # 额外模式：手机号
                        for m in re.finditer(r'1[3-9]\d{9}', remaining):
                            phone = m.group(0)
                            if not any(o == phone for o, _ in identified_originals) and phone != original:
                                type_counters[category] = type_counters.get(category, 0) + 1
                                rep = rule.get('replacement', '【电话${index}】').replace('${index}', str(type_counters[category]))
                                redactions.append({'original': phone, 'replacement': rep, 'type': category, 'location': ''})
                                identified_originals.append((phone, category))

                        # 额外模式：座机（带区号-分隔）
                        for m in re.finditer(r'0\d{2,3}-?\d{7,8}', remaining):
                            phone = m.group(0)
                            if not any(o == phone for o, _ in identified_originals) and phone != original:
                                type_counters[category] = type_counters.get(category, 0) + 1
                                rep = rule.get('replacement', '【电话${index}】').replace('${index}', str(type_counters[category]))
                                redactions.append({'original': phone, 'replacement': rep, 'type': category, 'location': ''})
                                identified_originals.append((phone, category))

                        # 额外模式：座机（括号区号）
                        for m in re.finditer(r'\(0\d{2,3}\)\d{7,8}', remaining):
                            phone = m.group(0)
                            if not any(o == phone for o, _ in identified_originals) and phone != original:
                                type_counters[category] = type_counters.get(category, 0) + 1
                                rep = rule.get('replacement', '【电话${index}】').replace('${index}', str(type_counters[category]))
                                redactions.append({'original': phone, 'replacement': rep, 'type': category, 'location': ''})
                                identified_originals.append((phone, category))

    # ---- 2. 黑名单处理（auto rules 之后）----
    for item in blacklist:
        original = item.get('original', '') if isinstance(item, dict) else str(item)
        item_type = item.get('type', 'blacklist') if isinstance(item, dict) else 'blacklist'
        type_name = TYPE_NAMES.get(item_type, TYPE_NAMES.get('blacklist', '黑名单'))

        if original and original in full_text and not any(o == original for o, _ in identified_originals):
            type_counters[item_type] = type_counters.get(item_type, 0) + 1
            replacement = f'【{type_name}{type_counters[item_type]}】'
            redactions.append({
                'original': original,
                'replacement': replacement,
                'type': item_type,
                'location': '',
            })
            identified_originals.append((original, item_type))

    # ---- 3. 自定义类型处理 ----
    for type_name, items in custom_types.items():
        for item in items:
            original = item.get('original', '')
            if not original:
                continue

            # 黑名单优先
            in_blacklist = any(b.get('original', '') == original for b in blacklist)
            if in_blacklist:
                continue

            # 白名单
            if is_covered_by_whitelist(original, whitelist):
                continue

            if original in full_text and not any(o == original for o, _ in identified_originals):
                ct_key = f'custom-{type_name}'
                type_counters[ct_key] = type_counters.get(ct_key, 0) + 1
                replacement = f'【{type_name}{type_counters[ct_key]}】'
                redactions.append({
                    'original': original,
                    'replacement': replacement,
                    'type': ct_key,
                    'location': '自定义类型',
                })
                identified_originals.append((original, ct_key))

    return redactions


# ============================================================
# DOCX 操作
# ============================================================

def replace_text_in_paragraph(para, old_text, new_text):
    """在段落中替换文本（处理跨 run 情况），保留原有格式"""
    if not old_text or len(old_text) < 2:
        return

    full_text = para.text
    if old_text not in full_text:
        return

    start_pos = full_text.find(old_text)
    if start_pos == -1:
        return
    end_pos = start_pos + len(old_text)

    # 找到包含目标文本的 runs
    char_count = 0
    target_runs = []

    for run in para.runs:
        run_start = char_count
        run_end = char_count + len(run.text)
        overlap_start = max(start_pos, run_start)
        overlap_end = min(end_pos, run_end)

        if overlap_start < overlap_end:
            target_runs.append({
                'run': run,
                'run_start': run_start,
                'run_end': run_end,
                'overlap_start': overlap_start,
                'overlap_end': overlap_end,
            })

        char_count = run_end

    if not target_runs:
        return

    if len(target_runs) == 1:
        target_runs[0]['run'].text = target_runs[0]['run'].text.replace(old_text, new_text)
    else:
        # 跨 run 替换
        first = target_runs[0]
        prefix = first['run'].text[:first['overlap_start'] - first['run_start']]
        first['run'].text = prefix + new_text

        for ri in target_runs[1:-1]:
            ri['run'].text = ''

        last = target_runs[-1]
        suffix = last['run'].text[last['overlap_end'] - last['run_start']:]
        last['run'].text = suffix


def apply_redactions_to_doc(doc, redactions):
    """将识别到的脱敏项应用到 DOCX 文档的所有段落和表格"""
    # 构建替换映射：original -> replacement
    # 按长度降序排列，避免短文本先替换破坏长文本
    sorted_redactions = sorted(redactions, key=lambda r: len(r['original']), reverse=True)

    # 处理正文段落
    for para in doc.paragraphs:
        for r in sorted_redactions:
            replace_text_in_paragraph(para, r['original'], r['replacement'])

    # 处理表格（递归嵌套）
    def process_tables(tables):
        for table in tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for r in sorted_redactions:
                            replace_text_in_paragraph(para, r['original'], r['replacement'])
                    # 嵌套表格
                    if cell.tables:
                        process_tables(cell.tables)

    if doc.tables:
        process_tables(doc.tables)


# ============================================================
# 文件操作
# ============================================================

def get_assets_dir():
    return Path(__file__).parent.parent / 'assets'


def load_user_settings():
    """加载用户设置"""
    settings_path = get_assets_dir() / 'user_settings.json'
    if not settings_path.exists():
        return {'whitelist': [], 'blacklist': [], 'customTypes': {}, 'disabledRules': []}
    try:
        with open(settings_path, 'r', encoding='utf-8') as f:
            return json.load(f)
    except (json.JSONDecodeError, IOError) as e:
        print(f"警告: 读取用户设置失败 - {e}")
        return {'whitelist': [], 'blacklist': [], 'customTypes': {}, 'disabledRules': []}


def generate_mapping_file(redactions, original_path, whitelist):
    """生成比对文件（Markdown 表格）"""
    # 基于原始文件名生成：xxx.docx → xxx_mapping.md（与 HTML 一致）
    base_name = os.path.splitext(os.path.basename(original_path))[0]
    dir_name = os.path.dirname(original_path)
    mapping_path = os.path.join(dir_name, f'{base_name}_mapping.md')
    filtered = [r for r in redactions if not is_covered_by_whitelist(r['original'], whitelist)]

    with open(mapping_path, 'w', encoding='utf-8') as f:
        f.write('# 脱敏内容替换比对\n\n')
        f.write(f'**文件**: {os.path.basename(original_path)}\n')
        f.write(f'**生成时间**: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}\n')
        f.write(f'**脱敏项数量**: {len(filtered)}\n\n')
        f.write('## 替换明细\n\n')
        f.write('| 序号 | 原文 | 替换 | 类型 |\n')
        f.write('|------|------|------|------|\n')
        for i, item in enumerate(filtered, 1):
            orig = item['original'].replace('|', '\\|')
            rep = item['replacement'].replace('|', '\\|')
            t = item['type']
            f.write(f'| {i} | `{orig}` | {rep} | {t} |\n')

    return mapping_path


def show_notification(title, message):
    """显示 Windows 通知"""
    try:
        import subprocess
        ps_cmd = (
            f'[Windows.UI.Notifications.ToastNotificationManager, Windows.UI.Notifications, '
            f'ContentType = WindowsRuntime] > $null; '
            f'$template = [Windows.UI.Notifications.ToastNotificationManager]::GetTemplateContent('
            f'[Windows.UI.Notifications.ToastTemplateType]::ToastText02); '
            f'$textNodes = $template.GetElementsByTagName("text"); '
            f'$textNodes.Item(0).Append($template.CreateTextNode("{title}")) > $null; '
            f'$textNodes.Item(1).Append($template.CreateTextNode("{message}")) > $null; '
            f'$toast = [Windows.UI.Notifications.ToastNotification]::new($template); '
            f'[Windows.UI.Notifications.ToastNotificationManager]::CreateToastNotifier("脱敏工具").Show($toast)'
        )
        subprocess.run(['powershell', '-Command', ps_cmd], capture_output=True, timeout=5)
    except Exception:
        pass


def main():
    if len(sys.argv) < 2:
        print("用法: python quick_redact.py <docx文件路径>")
        sys.exit(1)

    file_path = os.path.abspath(sys.argv[1])
    if not os.path.exists(file_path):
        print(f"错误: 文件不存在 - {file_path}")
        sys.exit(1)
    if not file_path.lower().endswith('.docx'):
        print(f"错误: 仅支持 .docx 文件 - {file_path}")
        sys.exit(1)

    # 输出路径
    dir_name = os.path.dirname(file_path)
    base_name = os.path.splitext(os.path.basename(file_path))[0]
    output_file = os.path.join(dir_name, f'{base_name}_脱敏.docx')

    # 加载用户设置
    settings = load_user_settings()
    whitelist = set(settings.get('whitelist', []))
    blacklist = settings.get('blacklist', [])
    custom_types = settings.get('customTypes', {})
    disabled_rules = set(settings.get('disabledRules', []))

    if whitelist:
        print(f"已加载白名单: {len(whitelist)} 项")
    if blacklist:
        print(f"已加载黑名单: {len(blacklist)} 项")
    if custom_types:
        print(f"已加载自定义类型: {len(custom_types)} 种")
    if disabled_rules:
        print(f"已禁用规则: {len(disabled_rules)} 个")

    # 读取 DOCX，拼接全文
    print(f"正在处理: {os.path.basename(file_path)}")
    doc = Document(file_path)

    # 收集所有段落文本（正文 + 表格）
    all_paragraphs = []

    def collect_paragraphs(paras):
        for para in paras:
            all_paragraphs.append(para.text)

    def collect_tables(tables):
        for table in tables:
            for row in table.rows:
                for cell in row.cells:
                    collect_paragraphs(cell.paragraphs)
                    if cell.tables:
                        collect_tables(cell.tables)

    collect_paragraphs(doc.paragraphs)
    if doc.tables:
        collect_tables(doc.tables)

    full_text = '\n'.join(all_paragraphs)

    # 执行识别
    redactions = identify_redactions(full_text, disabled_rules, whitelist, blacklist, custom_types)

    # 应用替换
    apply_redactions_to_doc(doc, redactions)

    # 保存脱敏文件
    doc.save(output_file)

    # 生成比对文件
    mapping_path = generate_mapping_file(redactions, file_path, whitelist)

    # 显示结果
    total = len(redactions)
    print(f"\n处理完成！")
    print(f"  脱敏文件: {output_file}")
    print(f"  比对文件: {mapping_path}")
    print(f"  共处理: {total} 处脱敏")

    show_notification('脱敏完成', f'{os.path.basename(file_path)} - 共 {total} 处脱敏')


if __name__ == '__main__':
    main()
