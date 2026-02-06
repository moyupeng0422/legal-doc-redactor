#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
法律文件脱敏脚本

功能：
1. 读取docx文件
2. 应用脱敏规则JSON
3. 生成脱敏后的docx文件
4. 保留原文格式（字体、段落、表格）

用法：
python redact.py input.docx rules.json -o output.docx
"""

import sys
import json
import argparse
import re
from datetime import datetime
from docx import Document
from docx.shared import RGBColor
from docx.oxml.ns import qn
from docx.enum.text import WD_COLOR_INDEX
# docxedit import removed - using native python-docx methods instead


class DocxRedactor:
    """docx文件脱敏处理器"""

    def __init__(self, rules_file, debug=False):
        """初始化脱敏器

        Args:
            rules_file: 脱敏规则JSON文件路径
            debug: 是否启用调试输出
        """
        self.debug = debug
        self.rules_data = self.load_rules(rules_file)
        self.rules = self._parse_rules()
        self.replacement_map = {}  # 原文 → 替换文本的映射
        self.redaction_log = []     # 脱敏操作日志

    def load_rules(self, rules_file):
        """加载脱敏规则"""
        try:
            with open(rules_file, 'r', encoding='utf-8') as f:
                return json.load(f)
        except FileNotFoundError:
            print(f"错误: 规则文件 {rules_file} 不存在")
            sys.exit(1)
        except json.JSONDecodeError as e:
            print(f"错误: 规则文件JSON格式错误 - {e}")
            sys.exit(1)

    def _parse_rules(self):
        """解析规则，返回按优先级排序的规则列表"""
        redactions = []

        # 检查是否有直接定义的redactions（前端导出格式）
        if 'redactions' in self.rules_data:
            return self.rules_data['redactions']

        # 否则使用rules配置
        rules = self.rules_data.get('rules', [])
        # 按优先级排序（降序）
        sorted_rules = sorted(rules, key=lambda x: x.get('priority', 0), reverse=True)

        # 转换为便于处理的格式
        for rule in sorted_rules:
            if not rule.get('enabled', True):
                continue

            patterns = rule.get('patterns', [])
            for pattern in patterns:
                try:
                    regex = re.compile(pattern)
                    redactions.append({
                        'regex': regex,
                        'rule': rule
                    })
                except re.error as e:
                    print(f"警告: 正则表达式错误 {pattern} - {e}")

        return redactions

    def apply_redaction(self, input_file, output_file):
        """应用脱敏到docx文件

        Args:
            input_file: 输入docx文件路径
            output_file: 输出docx文件路径
        """
        try:
            doc = Document(input_file)
        except Exception as e:
            print(f"错误: 无法读取文件 {input_file} - {e}")
            sys.exit(1)

        # 检查是否为前端导出规则（包含original字段）
        has_direct_redactions = any('original' in r for r in self.rules)

        if has_direct_redactions:
            # 前端规则：使用简化替换（保留格式，无高亮）
            self._simple_replace_with_format(doc)
        else:
            # 正则规则：使用原有复杂逻辑（保留格式+黄色高亮）
            # 处理段落
            for para_idx, para in enumerate(doc.paragraphs):
                self._redact_paragraph(para, f"paragraph:{para_idx + 1}")

            # 处理表格（包括嵌套表格）
            self._process_all_tables(doc)

            # 处理页眉页脚
            for section_idx, section in enumerate(doc.sections):
                # 主页眉
                if section.header:
                    for para_idx, para in enumerate(section.header.paragraphs):
                        self._redact_paragraph(para, f"header:section:{section_idx + 1},para:{para_idx + 1}")

                # 主页脚
                if section.footer:
                    for para_idx, para in enumerate(section.footer.paragraphs):
                        self._redact_paragraph(para, f"footer:section:{section_idx + 1},para:{para_idx + 1}")

                # 首页页眉（如果不同）
                if hasattr(section, 'first_page_header') and section.first_page_header:
                    for para_idx, para in enumerate(section.first_page_header.paragraphs):
                        self._redact_paragraph(para, f"first_header:section:{section_idx + 1},para:{para_idx + 1}")

                # 首页页脚（如果不同）
                if hasattr(section, 'first_page_footer') and section.first_page_footer:
                    for para_idx, para in enumerate(section.first_page_footer.paragraphs):
                        self._redact_paragraph(para, f"first_footer:section:{section_idx + 1},para:{para_idx + 1}")

                # 奇数页页眉（如果不同）
                if hasattr(section, 'even_page_header') and section.even_page_header:
                    for para_idx, para in enumerate(section.even_page_header.paragraphs):
                        self._redact_paragraph(para, f"even_header:section:{section_idx + 1},para:{para_idx + 1}")

                # 偶数页页脚（如果不同）
                if hasattr(section, 'even_page_footer') and section.even_page_footer:
                    for para_idx, para in enumerate(section.even_page_footer.paragraphs):
                        self._redact_paragraph(para, f"even_footer:section:{section_idx + 1},para:{para_idx + 1}")

        # 保存文档
        try:
            doc.save(output_file)
            print(f"成功: 脱敏文件已保存到 {output_file}")
            print(f"共处理 {len(self.redaction_log)} 处脱敏")
        except Exception as e:
            print(f"错误: 无法保存文件 {output_file} - {e}")
            sys.exit(1)

        return self.redaction_log

    def _simple_replace_with_format(self, doc):
        """使用简化替换（保留段落格式，不添加黄色高亮）

        用于前端导出的规则，原始文本已去除自动编号前缀

        Args:
            doc: python-docx文档对象
        """
        # 处理正文段落
        for para in doc.paragraphs:
            self._replace_in_para_simple(para)

        # 处理表格（包括嵌套表格）
        for table in doc.tables:
            self._process_table_simple(table)

        # 处理页眉页脚
        for section_idx, section in enumerate(doc.sections):
            # 主页眉
            if section.header:
                for para in section.header.paragraphs:
                    self._replace_in_para_simple(para)

            # 主页脚
            if section.footer:
                for para in section.footer.paragraphs:
                    self._replace_in_para_simple(para)

            # 首页页眉（如果不同）
            if hasattr(section, 'first_page_header') and section.first_page_header:
                for para in section.first_page_header.paragraphs:
                    self._replace_in_para_simple(para)

            # 首页页脚（如果不同）
            if hasattr(section, 'first_page_footer') and section.first_page_footer:
                for para in section.first_page_footer.paragraphs:
                    self._replace_in_para_simple(para)

            # 奇数页页眉（如果不同）
            if hasattr(section, 'even_page_header') and section.even_page_header:
                for para in section.even_page_header.paragraphs:
                    self._replace_in_para_simple(para)

            # 偶数页页脚（如果不同）
            if hasattr(section, 'even_page_footer') and section.even_page_footer:
                for para in section.even_page_footer.paragraphs:
                    self._replace_in_para_simple(para)

    def _process_table_simple(self, table):
        """递归处理表格及其嵌套表格（简化版本）"""
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    self._replace_in_para_simple(para)
                # 递归处理嵌套表格
                for nested_table in cell.tables:
                    self._process_table_simple(nested_table)

    def _replace_in_para_simple(self, para):
        """在段落中进行简单替换（保留格式，不添加高亮）

        使用混合方式：先尝试简单替换，失败则使用跨runs处理

        Args:
            para: python-docx段落对象
        """
        for redaction in self.rules:
            if 'original' in redaction:
                original = redaction['original']
                replacement = redaction['replacement']

                # 验证长度（人名最小长度为2，其他类型为4）
                min_length = 2 if redaction.get('type') == 'person_name' else 4
                if not original or len(original.strip()) < min_length:
                    if self.debug:
                        print(f"[DEBUG]   跳过（长度<{min_length}）: '{original}'")
                    continue

                # 检查是否在段落中存在
                if original not in para.text:
                    continue

                # 策略1：先尝试简单替换（单run内）
                replaced = False
                for run in para.runs:
                    if original in run.text:
                        run.text = run.text.replace(original, replacement)
                        replaced = True
                        break

                if replaced:
                    if self.debug:
                        print(f"[DEBUG]   [OK] 简单替换: '{original[:30]}...' -> '{replacement[:30]}...'")
                    self.redaction_log.append(redaction)
                    continue

                # 策略2：处理跨runs情况
                # 在runs中累计查找目标文本的位置
                char_count = 0
                start_pos = -1
                for run in para.runs:
                    pos_in_run = run.text.find(original)
                    if pos_in_run != -1:
                        start_pos = char_count + pos_in_run
                        break
                    char_count += len(run.text)

                if start_pos == -1:
                    # 跨runs情况：需要拼接runs来查找
                    char_count = 0
                    accumulated_text = ""
                    target_runs = []
                    end_pos = -1

                    for run in para.runs:
                        run_start = char_count
                        accumulated_text += run.text
                        run_end = char_count + len(run.text)

                        # 在累积文本中查找
                        if original in accumulated_text:
                            found_pos = accumulated_text.find(original)
                            start_pos = found_pos
                            end_pos = start_pos + len(original)

                            # 找到包含目标文本的runs
                            temp_count = 0
                            for r in para.runs:
                                r_start = temp_count
                                r_end = temp_count + len(r.text)
                                overlap_start = max(start_pos, r_start)
                                overlap_end = min(end_pos, r_end)
                                if overlap_start < overlap_end:
                                    target_runs.append({
                                        'run': r,
                                        'run_start': r_start,
                                        'run_end': r_end,
                                        'overlap_start': overlap_start,
                                        'overlap_end': overlap_end
                                    })
                                temp_count = r_end
                            break

                        char_count = run_end

                    if not target_runs:
                        if self.debug:
                            print(f"[DEBUG]   跳过（跨runs查找失败）: '{original[:50]}...'")
                        continue

                    # 执行跨runs替换
                    if len(target_runs) == 1:
                        target_runs[0]['run'].text = target_runs[0]['run'].text.replace(original, replacement)
                    else:
                        first_run = target_runs[0]['run']
                        overlap_start = target_runs[0]['overlap_start']
                        overlap_end = target_runs[0]['run_end']
                        prefix = first_run.text[:overlap_start - target_runs[0]['run_start']]

                        last_run = target_runs[-1]['run']
                        overlap_start = target_runs[-1]['overlap_start']
                        overlap_end = target_runs[-1]['overlap_end']
                        suffix = last_run.text[overlap_end - target_runs[-1]['run_start']:]

                        first_run.text = prefix + replacement

                        for run_info in target_runs[1:-1]:
                            run_info['run'].text = ''

                        last_run.text = suffix

                    if self.debug:
                        print(f"[DEBUG]   [OK] 跨runs替换: '{original[:30]}...' -> '{replacement[:30]}...'")
                    self.redaction_log.append(redaction)

    def _redact_paragraph(self, para, location):
        """对单个段落应用脱敏

        Args:
            para: python-docx段落对象
            location: 位置标识
        """
        # 获取段落文本
        original_text = para.text
        # 移除空文本检查 - 自动标号段落可能para.text为空，但仍需处理
        # if not original_text or not original_text.strip():
        #     return

        # 应用所有规则
        redactions_in_para = []

        # 第一步：收集前端导出的直接替换项
        for redaction in self.rules:
            if 'original' in redaction:
                original = redaction.get('original', '')
                # 验证长度（避免空/短文本）
                if original and len(original.strip()) >= 4 and original in para.text:
                    redactions_in_para.append({
                        'original': original,
                        'replacement': redaction.get('replacement', ''),
                        'type': redaction.get('type', 'unknown'),
                        'location': location
                    })

        # 第二步：应用正则匹配规则
        for redaction in self.rules:
            # 跳过已处理的直接替换项
            if 'original' in redaction:
                continue

            regex = redaction['regex']
            rule = redaction['rule']

            for match in regex.finditer(original_text):
                # 检查是否使用捕获组（useCaptureGroup）
                # 如果规则设置了 useCaptureGroup=true 且存在捕获组，则使用第一个捕获组
                if rule.get('useCaptureGroup', False) and match.lastindex and match.lastindex >= 1:
                    matched_text = match.group(1)  # 使用捕获组内容
                else:
                    matched_text = match.group(0)  # 使用完整匹配
                replacement = self._generate_replacement(matched_text, rule, len(redactions_in_para) + 1)

                redactions_in_para.append({
                    'original': matched_text,
                    'replacement': replacement,
                    'type': rule.get('category', 'unknown'),
                    'location': location
                })

        # 第三步：按长度排序，统一应用替换
        redactions_in_para.sort(key=lambda r: len(r['original']), reverse=True)
        for redaction in redactions_in_para:
            self._replace_in_paragraph(para, redaction)
            self.redaction_log.append(redaction)

    def _apply_direct_redaction(self, para, redaction, location):
        """应用前端导出的直接替换

        Args:
            para: 段落对象
            redaction: 脱敏项（包含original和replacement）
            location: 位置标识
        """
        original = redaction.get('original', '')
        replacement = redaction.get('replacement', '')

        if original in para.text:
            self._replace_text_in_paragraph(para, original, replacement)
            self.redaction_log.append({
                'original': original,
                'replacement': replacement,
                'type': redaction.get('type', 'unknown'),
                'location': location
            })

    def _generate_replacement(self, matched_text, rule, index):
        """生成替换文本

        Args:
            matched_text: 匹配的原文
            rule: 规则配置
            index: 同类序号

        Returns:
            替换文本
        """
        replacement_template = rule.get('replacement', '【脱敏内容】')

        # 简单替换（不支持上下文）
        replacement = replacement_template
        replacement = replacement.replace('${index}', str(index))
        replacement = replacement.replace('${category}', rule.get('category', ''))

        return replacement

    def _replace_in_paragraph(self, para, redaction):
        """在段落中执行替换

        Args:
            para: 段落对象
            redaction: 脱敏项
        """
        original = redaction['original']
        replacement = redaction['replacement']
        self._replace_text_in_paragraph(para, original, replacement)

    def _replace_text_in_paragraph(self, para, old_text, new_text):
        """在段落中替换文本并添加高亮（保留原有格式）

        处理跨 run 的文本替换，确保格式正确保留

        Args:
            para: 段落对象
            old_text: 要替换的文本
            new_text: 新文本
        """
        # 验证：old_text 不能为空或太短（地址通常至少4个字符）
        if not old_text or len(old_text.strip()) < 4:
            return

        # 获取段落完整文本
        full_text = para.text

        # 检查 old_text 是否在段落中
        if old_text not in full_text:
            return

        # 找到 old_text 的位置
        start_pos = full_text.find(old_text)
        if start_pos == -1:
            return

        end_pos = start_pos + len(old_text)

        # 遍历 runs，找到包含目标文本的 runs
        char_count = 0
        target_runs = []

        for run in para.runs:
            run_text = run.text
            run_start = char_count
            run_end = char_count + len(run_text)

            # 检查这个 run 是否与目标文本重叠
            overlap_start = max(start_pos, run_start)
            overlap_end = min(end_pos, run_end)

            if overlap_start < overlap_end:
                # 这个 run 包含目标文本的一部分
                target_runs.append({
                    'run': run,
                    'run_start': run_start,
                    'run_end': run_end,
                    'overlap_start': overlap_start,
                    'overlap_end': overlap_end
                })

            char_count = run_end

        if not target_runs:
            return

        # 处理替换
        if len(target_runs) == 1:
            # 简单情况：目标文本在一个 run 中
            run_info = target_runs[0]
            run = run_info['run']
            run.text = run.text.replace(old_text, new_text)
            run.font.highlight_color = WD_COLOR_INDEX.YELLOW
        else:
            # 复杂情况：目标文本跨越多个 runs
            # 将第一个 run 中目标文本前的部分保留
            first_run = target_runs[0]['run']
            overlap_start = target_runs[0]['overlap_start']
            overlap_end = target_runs[0]['run_end']
            prefix = first_run.text[:overlap_start - target_runs[0]['run_start']]

            # 将最后一个 run 中目标文本后的部分保留
            last_run = target_runs[-1]['run']
            overlap_start = target_runs[-1]['overlap_start']
            overlap_end = target_runs[-1]['overlap_end']
            suffix = last_run.text[overlap_end - target_runs[-1]['run_start']:]

            # 修改第一个 run：保留前缀 + 新文本
            first_run.text = prefix + new_text
            first_run.font.highlight_color = WD_COLOR_INDEX.YELLOW

            # 清空中间的 runs
            for run_info in target_runs[1:-1]:
                run_info['run'].text = ''

            # 修改最后一个 run：只保留后缀
            last_run.text = suffix

    def _process_all_tables(self, doc):
        """递归处理所有表格（包括嵌套表格）"""
        for table_idx, table in enumerate(doc.tables):
            self._process_table_recursive(table, f"table:{table_idx + 1}")

    def _process_table_recursive(self, table, location_prefix):
        """递归处理表格及其嵌套表格"""
        for row_idx, row in enumerate(table.rows):
            for cell_idx, cell in enumerate(row.cells):
                # 处理单元格段落
                for para_idx, para in enumerate(cell.paragraphs):
                    location = f"{location_prefix},row:{row_idx + 1},cell:{cell_idx + 1},para:{para_idx + 1}"
                    self._redact_paragraph(para, location)

                # 递归处理嵌套表格
                for nested_table_idx, nested_table in enumerate(cell.tables):
                    nested_location = f"{location_prefix},row:{row_idx + 1},cell:{cell_idx + 1},nested_table:{nested_table_idx + 1}"
                    self._process_table_recursive(nested_table, nested_location)


def main():
    """主函数"""
    parser = argparse.ArgumentParser(description='法律文件脱敏工具')
    parser.add_argument('input_file', help='输入docx文件')
    parser.add_argument('rules_file', help='脱敏规则JSON文件')
    parser.add_argument('-o', '--output', help='输出docx文件（默认：input_脱敏.docx）')
    parser.add_argument('--log', help='输出处理日志JSON文件')
    parser.add_argument('--debug', action='store_true', help='启用调试输出（显示空段落和runs信息）')

    args = parser.parse_args()

    # 确定输出文件名
    if args.output:
        output_file = args.output
    else:
        input_name = args.input_file
        if input_name.endswith('.docx'):
            output_file = input_name[:-5] + '_脱敏.docx'
        else:
            output_file = input_name + '_脱敏.docx'

    # 创建脱敏器
    redactor = DocxRedactor(args.rules_file, debug=args.debug)

    # 执行脱敏
    print(f"正在处理: {args.input_file}")
    redaction_log = redactor.apply_redaction(args.input_file, output_file)

    # 输出日志
    if args.log:
        log_data = {
            'inputFile': args.input_file,
            'outputFile': output_file,
            'processTime': datetime.now().isoformat(),
            'totalRedactions': len(redaction_log),
            'redactions': redaction_log
        }
        with open(args.log, 'w', encoding='utf-8') as f:
            json.dump(log_data, f, ensure_ascii=False, indent=2)
        print(f"日志已保存到: {args.log}")


if __name__ == '__main__':
    main()
