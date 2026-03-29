#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
法律文件脱敏处理 - 一键还原脚本

完整复刻 HTML 应用 restoreSingleFile() 的还原逻辑，
自动查找比对文件，将脱敏稿还原为原文。

用法：python quick_restore.py <脱敏稿.docx>
"""

import os
import sys
import re
import subprocess
from datetime import datetime

try:
    from docx import Document
except ImportError:
    print("错误: 缺少 python-docx 库，请运行: pip install python-docx")
    sys.exit(1)


# ============================================================
# 映射文件解析（完整复刻 HTML parseMapping）
# ============================================================

def parse_mapping(content):
    """解析比对 markdown 文件，提取替换映射

    与 HTML parseMapping() 完全一致：
    - 同一正则解析 markdown 表格
    - 同一清理规则（反引号、链接、尖括号）
    - 同一过滤（只处理 【...】 标记）

    Args:
        content: 比对文件内容

    Returns:
        dict {replacement_marker: original_text}
    """
    mapping = {}
    table_pattern = re.compile(r'\|\s*\d+\s*\|\s*([^|]+)\s*\|\s*([^|]+)\s*\|', re.MULTILINE)

    for match in table_pattern.finditer(content):
        original = match.group(1).strip()
        replacement = match.group(2).strip()

        # 跳过表头
        if original == '原文' or replacement == '替换':
            continue

        # 只处理 【...】 格式的替换标记
        if replacement.startswith('【') and replacement.endswith('】'):
            # 清理 markdown 格式（与 HTML 一致）
            # 1. 代码块格式: `xxx@example.com`
            code_block_match = re.match(r'^`(.+)`$', original)
            if code_block_match:
                original = code_block_match.group(1)
            # 2. markdown链接格式: [xxx@example.com](mailto:xxx@example.com)
            elif original.startswith('[') and '](' in original:
                link_text_match = re.match(r'^\[([^\]]+)\]', original)
                if link_text_match:
                    original = link_text_match.group(1)
            # 3. 自动链接格式: <xxx@example.com>
            elif original.startswith('<') and original.endswith('>'):
                original = original[1:-1]

            mapping[replacement] = original

    return mapping


# ============================================================
# 文本替换逻辑（完整复刻 HTML restoreSingleFile）
# ============================================================

def build_run_info(para):
    """构建段落的 run 信息（与 HTML 累积文本+位置记录一致）

    Args:
        para: python-docx 段落对象

    Returns:
        (accumulated_text, run_info_list)
    """
    accumulated_text = ""
    run_info = []

    for i, run in enumerate(para.runs):
        run_text = run.text if run.text else ""
        start_pos = len(accumulated_text)
        accumulated_text += run_text
        end_pos = len(accumulated_text)

        run_info.append({
            'run': run,
            'start_pos': start_pos,
            'end_pos': end_pos,
            'text': run_text,
        })

    return accumulated_text, run_info


def restore_paragraph(para, mapping):
    """还原段落中的脱敏内容（完整复刻 HTML 逻辑）

    关键行为（与 HTML 一致）：
    1. 用 str.find() 纯字符串匹配（非正则）
    2. 每次替换后立即重建 runInfo（HTML 的关键行为）
    3. 跨 run 处理：首 run 保留前缀+原文，中间 run 清空，末 run 保留后缀

    Args:
        para: python-docx 段落对象
        mapping: {replacement: original} 字典

    Returns:
        还原的数量
    """
    if not mapping:
        return 0

    count = 0

    # 按替换标记长度降序排列（与 HTML 一致：优先替换长文本）
    sorted_items = sorted(mapping.items(), key=lambda x: len(x[0]), reverse=True)

    for replacement, original in sorted_items:
        # 构建当前段落的位置信息
        accumulated_text, run_info = build_run_info(para)

        # 用 str.find() 纯字符串匹配（与 HTML indexOf 一致）
        pos = accumulated_text.find(replacement)

        while pos != -1:
            end_pos = pos + len(replacement)

            # 找到受影响的 runs（与 HTML 重叠判断一致）
            affected_runs = []
            for info in run_info:
                overlap_start = max(pos, info['start_pos'])
                overlap_end = min(end_pos, info['end_pos'])
                # 与 HTML 完全一致的重叠判断
                has_overlap = (overlap_start <= overlap_end
                              and overlap_start < info['end_pos']
                              and overlap_end > info['start_pos'])
                if has_overlap:
                    affected_runs.append({
                        'info': info,
                        'overlap_start': overlap_start,
                        'overlap_end': overlap_end,
                    })

            if not affected_runs:
                break

            if len(affected_runs) == 1:
                # 单 run 内替换（与 HTML 一致）
                info = affected_runs[0]['info']
                start_offset = affected_runs[0]['overlap_start'] - info['start_pos']
                end_offset = affected_runs[0]['overlap_end'] - info['start_pos']
                new_content = info['text'][:start_offset] + original + info['text'][end_offset:]
                info['run'].text = new_content
            else:
                # 跨 run 替换（与 HTML 一致）
                first_info = affected_runs[0]['info']
                last_info = affected_runs[-1]['info']

                first_start = affected_runs[0]['overlap_start'] - first_info['start_pos']
                before = first_info['text'][:first_start]

                last_end = affected_runs[-1]['overlap_end'] - last_info['start_pos']
                after = last_info['text'][last_end:]

                # 首 run：保留前缀 + 原文
                first_info['run'].text = before + original
                # 中间 run：清空
                for item in affected_runs[1:-1]:
                    item['info']['run'].text = ""
                # 末 run：保留后缀
                last_info['run'].text = after

            count += 1

            # 关键：每次替换后立即重建 runInfo（与 HTML 一致）
            accumulated_text, run_info = build_run_info(para)

            # 从 pos+1 继续搜索（与 HTML batch 版本一致）
            pos = accumulated_text.find(replacement, pos + 1)

    return count


# ============================================================
# 还原主流程
# ============================================================

def restore_doc(doc, mapping):
    """将还原应用到 DOCX 文档的所有段落和表格

    Args:
        doc: python-docx Document 对象
        mapping: {replacement: original} 字典

    Returns:
        还原的总数量
    """
    restore_count = 0

    # 处理正文段落
    for para in doc.paragraphs:
        count = restore_paragraph(para, mapping)
        restore_count += count

    # 处理表格（含嵌套表格）
    def process_tables(tables):
        nonlocal restore_count
        for table in tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        count = restore_paragraph(para, mapping)
                        restore_count += count
                    # 嵌套表格
                    if cell.tables:
                        process_tables(cell.tables)

    if doc.tables:
        process_tables(doc.tables)

    return restore_count


# ============================================================
# 文件查找
# ============================================================

def find_mapping_file(docx_path):
    """自动查找对应的比对文件

    查找优先级（与 HTML autoMatchPairs 一致）：
    1. {base}_比对.md（HTML 批量模式生成）
    2. {base}_mapping.md（Python 一键脱敏生成）
    3. {base}比对.md
    4. {base}.md

    Args:
        docx_path: 脱敏稿 docx 文件路径

    Returns:
        比对文件路径，找不到则返回 None
    """
    dir_name = os.path.dirname(docx_path)
    # 去掉 【xxx脱敏】 后缀（与 HTML 一致）和 _脱敏 后缀（Python 一键脱敏生成）
    base_name = os.path.splitext(os.path.basename(docx_path))[0]
    base_name = re.sub(r'【[^】]*脱敏】$', '', base_name)
    base_name = re.sub(r'_脱敏$', '', base_name)

    candidates = [
        os.path.join(dir_name, f'{base_name}_比对.md'),
        os.path.join(dir_name, f'{base_name}_mapping.md'),
        os.path.join(dir_name, f'{base_name}比对.md'),
        os.path.join(dir_name, f'{base_name}.md'),
    ]

    for path in candidates:
        if os.path.exists(path):
            return path

    return None


# ============================================================
# 通知
# ============================================================

def show_notification(title, message):
    """显示 Windows 通知"""
    try:
        ps_cmd = (
            f'[Windows.UI.Notifications.ToastNotificationManager, Windows.UI.Notifications, '
            f'ContentType = WindowsRuntime] > $null; '
            f'$template = [Windows.UI.Notifications.ToastNotificationManager]::GetTemplateContent('
            f'[Windows.UI.Notifications.ToastTemplateType]::ToastText02); '
            f'$textNodes = $template.GetElementsByTagName("text"); '
            f'$textNodes.Item(0).Append($template.CreateTextNode("{title}")) > $null; '
            f'$textNodes.Item(1).Append($template.CreateTextNode("{message}")) > $null; '
            f'$toast = [Windows.UI.Notifications.ToastNotification]::new($template); '
            f'[Windows.UI.Notifications.ToastNotificationManager]::CreateToastNotifier("脱敏工具").Show($toast)'
        )
        subprocess.run(['powershell', '-Command', ps_cmd], capture_output=True, timeout=5)
    except Exception:
        pass


# ============================================================
# 主函数
# ============================================================

def main():
    if len(sys.argv) < 2:
        print("用法: python quick_restore.py <脱敏稿.docx>")
        sys.exit(1)

    file_path = os.path.abspath(sys.argv[1])
    if not os.path.exists(file_path):
        print(f"错误: 文件不存在 - {file_path}")
        sys.exit(1)
    if not file_path.lower().endswith('.docx'):
        print(f"错误: 仅支持 .docx 文件 - {file_path}")
        sys.exit(1)

    # 自动查找比对文件
    mapping_path = find_mapping_file(file_path)
    if not mapping_path:
        print(f"错误: 未找到比对文件")
        print(f"  已查找:")
        dir_name = os.path.dirname(file_path)
        base_name = re.sub(r'【[^】]*脱敏】', '', os.path.splitext(os.path.basename(file_path))[0])
        for name in [f'{base_name}_比对.md', f'{base_name}_mapping.md', f'{base_name}比对.md', f'{base_name}.md']:
            print(f"    {os.path.join(dir_name, name)}")
        show_notification('还原失败', '未找到比对文件')
        sys.exit(1)

    print(f"比对文件: {os.path.basename(mapping_path)}")

    # 解析映射
    with open(mapping_path, 'r', encoding='utf-8') as f:
        content = f.read()
    mapping = parse_mapping(content)

    if not mapping:
        print(f"错误: 比对文件中未解析到有效的替换映射")
        show_notification('还原失败', '比对文件为空')
        sys.exit(1)

    print(f"解析到 {len(mapping)} 条替换映射")

    # 输出文件名（与 HTML 一致：去掉 【xxx脱敏】 或 _脱敏 后缀，加 _还原）
    base_name = os.path.splitext(os.path.basename(file_path))[0]
    base_name = re.sub(r'【[^】]*脱敏】$', '', base_name)
    base_name = re.sub(r'_脱敏$', '', base_name)
    dir_name = os.path.dirname(file_path)
    output_file = os.path.join(dir_name, f'{base_name}_还原.docx')

    # 读取 DOCX 并还原
    print(f"正在处理: {os.path.basename(file_path)}")
    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"错误: 无法读取文件 - {e}")
        show_notification('还原失败', str(e))
        sys.exit(1)

    restore_count = restore_doc(doc, mapping)

    # 保存
    try:
        doc.save(output_file)
    except Exception as e:
        print(f"错误: 无法保存文件 - {e}")
        show_notification('还原失败', str(e))
        sys.exit(1)

    print(f"\n处理完成！")
    print(f"  还原文件: {output_file}")
    print(f"  共还原: {restore_count} 处内容")

    show_notification('还原完成', f'{os.path.basename(file_path)} - 共 {restore_count} 处还原')


if __name__ == '__main__':
    main()
