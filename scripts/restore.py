#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
法律文件还原脚本

功能：
1. 读取脱敏后的docx文件
2. 读取比对.md文档
3. 将脱敏内容还原为原始内容

用法：
python restore.py redacted.docx mapping.md -o restored.docx
"""

import sys
import re
import argparse
from datetime import datetime
from docx import Document


class DocxRestorer:
    """docx文件还原处理器"""

    def __init__(self, mapping_file):
        """初始化还原器

        Args:
            mapping_file: 替换比对.md文件路径
        """
        self.mapping = self.parse_mapping_file(mapping_file)

    def parse_mapping_file(self, mapping_file):
        """解析比对.md文档，提取替换映射

        Args:
            mapping_file: 比对.md文件路径

        Returns:
            替换映射字典 {替换文本: 原文}
        """
        mapping = {}

        try:
            with open(mapping_file, 'r', encoding='utf-8') as f:
                content = f.read()

            # 解析表格行
            # 格式: | 序号 | 原文 | 替换 | 类型 | 位置 |
            table_pattern = re.compile(r'\|\s*\d+\s*\|\s*([^|]+)\s*\|\s*([^|]+)\s*\|', re.MULTILINE)

            for match in table_pattern.finditer(content):
                original = match.group(1).strip()
                replacement = match.group(2).strip()

                # 跳过表头
                if original == '原文' or replacement == '替换':
                    continue

                # 检查是否为有效的脱敏替换
                if replacement.startswith('【') and replacement.endswith('】'):
                    # 处理各种markdown链接格式，提取真实的邮箱地址
                    # 1. 代码块格式: `xxx@example.com`
                    code_block_match = re.match(r'^`(.+)`$', original)
                    if code_block_match:
                        original = code_block_match.group(1)
                    # 2. markdown链接格式: [xxx@example.com](mailto:xxx@example.com)
                    elif original.startswith('[') and '](' in original:
                        link_text_match = re.match(r'^\[([^\]]+)\]', original)
                        if link_text_match:
                            original = link_text_match.group(1)
                    # 3. 自动链接格式: <xxx@example.com>
                    elif original.startswith('<') and original.endswith('>'):
                        original = original[1:-1]

                    mapping[replacement] = original

            print(f"从比对文档中读取了 {len(mapping)} 条替换映射")
            return mapping

        except FileNotFoundError:
            print(f"错误: 比对文档 {mapping_file} 不存在")
            sys.exit(1)
        except Exception as e:
            print(f"错误: 解析比对文档失败 - {e}")
            sys.exit(1)

    def restore(self, input_file, output_file):
        """还原docx文件

        Args:
            input_file: 脱敏后的docx文件路径
            output_file: 还原后的输出文件路径
        """
        try:
            doc = Document(input_file)
        except Exception as e:
            print(f"错误: 无法读取文件 {input_file} - {e}")
            sys.exit(1)

        restore_count = 0

        # 处理段落
        for para in doc.paragraphs:
            count = self._restore_paragraph(para)
            restore_count += count

        # 处理表格
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        count = self._restore_paragraph(para)
                        restore_count += count

        # 保存文档
        try:
            doc.save(output_file)
            print(f"成功: 还原文件已保存到 {output_file}")
            print(f"共还原 {restore_count} 处内容")
        except Exception as e:
            print(f"错误: 无法保存文件 {output_file} - {e}")
            sys.exit(1)

        return restore_count

    def _restore_paragraph(self, para):
        """还原段落中的脱敏内容

        Args:
            para: python-docx段落对象

        Returns:
            还原的数量
        """
        text = para.text
        count = 0

        # 调试：输出段落文本
        if any(rep in text for rep in self.mapping.keys()):
            print(f"调试：段落包含替换标记")
            print(f"  段落文本: {text[:100]}...")

        # 按长度降序排序，优先替换长文本
        sorted_replacements = sorted(self.mapping.items(), key=lambda x: len(x[0]), reverse=True)

        for replacement, original in sorted_replacements:
            if replacement in text:
                print(f"调试：找到 '{replacement}' -> '{original}'")
                self._replace_text_in_paragraph(para, replacement, original)
                text = text.replace(replacement, original)
                count += 1

        return count

    def _replace_text_in_paragraph(self, para, old_text, new_text):
        """在段落中替换文本（保留格式、修订、批注）

        使用runs级别替换，保留段落中的所有格式、修订痕迹、批注等。

        Args:
            para: 段落对象
            old_text: 要替换的文本（如【邮箱1】）
            new_text: 新文本（如原邮箱）
        """
        # 收集所有runs的文本和位置信息
        accumulated_text = ""
        run_info = []

        for i, run in enumerate(para.runs):
            run_text = run.text if run.text else ""
            start_pos = len(accumulated_text)
            accumulated_text += run_text
            end_pos = len(accumulated_text)

            run_info.append({
                'run': run,
                'start_pos': start_pos,
                'end_pos': end_pos,
                'text': run_text
            })

        # 查找要替换的文本
        pos = accumulated_text.find(old_text)
        if pos == -1:
            return

        # 找到包含目标文本的runs
        target_end = pos + len(old_text)
        affected_runs = []

        for info in run_info:
            # 检查此run是否与目标范围重叠
            overlap_start = max(pos, info['start_pos'])
            overlap_end = min(target_end, info['end_pos'])

            if overlap_start < overlap_end:
                affected_runs.append({
                    'info': info,
                    'overlap_start': overlap_start,
                    'overlap_end': overlap_end
                })

        if not affected_runs:
            return

        # 执行替换
        if len(affected_runs) == 1:
            # 在同一个run内替换
            info = affected_runs[0]['info']
            run = info['run']
            start_offset = affected_runs[0]['overlap_start'] - info['start_pos']
            end_offset = affected_runs[0]['overlap_end'] - info['start_pos']

            run.text = run.text[:start_offset] + new_text + run.text[end_offset:]
        else:
            # 跨多个runs替换
            first_info = affected_runs[0]['info']
            last_info = affected_runs[-1]['info']

            first_start = affected_runs[0]['overlap_start'] - first_info['start_pos']
            before = first_info['text'][:first_start]

            last_end = affected_runs[-1]['overlap_end'] - last_info['start_pos']
            after = last_info['text'][last_end:]

            # 第一个run：保留前缀 + 替换文本
            first_info['run'].text = before + new_text

            # 中间的runs：清空
            for item in affected_runs[1:-1]:
                item['info']['run'].text = ""

            # 最后一个run：保留后缀
            last_info['run'].text = after


def main():
    """主函数"""
    parser = argparse.ArgumentParser(description='法律文件还原工具')
    parser.add_argument('input_file', help='脱敏后的docx文件')
    parser.add_argument('mapping_file', help='替换比对.md文档')
    parser.add_argument('-o', '--output', help='输出docx文件（默认：input_还原.docx）')

    args = parser.parse_args()

    # 确定输出文件名
    if args.output:
        output_file = args.output
    else:
        input_name = args.input_file
        if input_name.endswith('.docx'):
            output_file = input_name[:-5] + '_还原.docx'
        else:
            output_file = input_name + '_还原.docx'

    # 创建还原器
    restorer = DocxRestorer(args.mapping_file)

    # 执行还原
    print(f"正在处理: {args.input_file}")
    restorer.restore(args.input_file, output_file)


if __name__ == '__main__':
    main()
